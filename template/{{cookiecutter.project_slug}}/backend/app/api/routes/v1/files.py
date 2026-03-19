{%- if cookiecutter.enable_ai_agent and cookiecutter.enable_conversation_persistence and cookiecutter.use_jwt %}
"""File upload and download endpoints for chat attachments."""

import logging
{%- if cookiecutter.use_postgresql %}
from uuid import UUID
{%- endif %}

from fastapi import APIRouter, UploadFile, File, HTTPException, status
from fastapi.responses import Response
from sqlalchemy import select

from app.api.deps import DBSession, CurrentUser
from app.db.models.chat_file import ChatFile
from app.schemas.file import FileUploadResponse, FileInfo
from app.services.file_storage import (
    get_file_storage,
    classify_file,
    ALLOWED_MIME_TYPES,
    MAX_UPLOAD_SIZE,
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/files", tags=["files"])


def _parse_text_content(data: bytes, mime_type: str) -> str | None:
    """Extract text content from text-based files."""
    try:
        return data.decode("utf-8")
    except (UnicodeDecodeError, ValueError):
        return None


{%- if not cookiecutter.use_llamaparse %}
def _parse_pdf_content(data: bytes) -> str | None:
    """Extract text from PDF using PyMuPDF."""
    try:
        import pymupdf
        doc = pymupdf.open(stream=data, filetype="pdf")
        texts = []
        for page in doc:
            blocks = page.get_text("blocks")
            for b in blocks:
                if b[6] == 0:
                    text = b[4].strip()
                    if text:
                        texts.append(text)
            try:
                tables = page.find_tables()
                if tables and tables.tables:
                    for table in tables.tables:
                        df = table.to_pandas()
                        if not df.empty:
                            texts.append(df.to_markdown(index=False))
            except Exception:
                pass
        doc.close()
        return "\n\n".join(texts) if texts else None
    except Exception as e:
        logger.warning(f"PDF parsing failed: {e}")
        return None


def _parse_docx_content(data: bytes) -> str | None:
    """Extract text from DOCX."""
    try:
        import io
        from docx import Document as DOCXDocument
        doc = DOCXDocument(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        logger.warning(f"DOCX parsing failed: {e}")
        return None
{%- endif %}


@router.post("/upload", response_model=FileUploadResponse, status_code=status.HTTP_201_CREATED)
async def upload_file(
    file: UploadFile = File(...),
    db: DBSession = None,
    current_user: CurrentUser = None,
):
    """Upload a file for use in chat."""
    if file.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"File type '{file.content_type}' is not supported.",
        )

    data = await file.read()
    if len(data) > MAX_UPLOAD_SIZE:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File too large. Maximum size is {MAX_UPLOAD_SIZE // (1024*1024)}MB.",
        )

    file_type = classify_file(file.content_type or "", file.filename or "unknown")

    parsed_content = None
    if file_type == "text":
        parsed_content = _parse_text_content(data, file.content_type or "")
{%- if not cookiecutter.use_llamaparse %}
    elif file_type == "pdf":
        parsed_content = _parse_pdf_content(data)
    elif file_type == "docx":
        parsed_content = _parse_docx_content(data)
{%- endif %}

    storage = get_file_storage()
    storage_path = await storage.save(str(current_user.id), file.filename or "unknown", data)

    chat_file = ChatFile(
        user_id=current_user.id,
        filename=file.filename or "unknown",
        mime_type=file.content_type or "application/octet-stream",
        size=len(data),
        storage_path=storage_path,
        file_type=file_type,
        parsed_content=parsed_content,
    )
    db.add(chat_file)
    await db.flush()
    await db.commit()
    await db.refresh(chat_file)

    return FileUploadResponse(
        id=chat_file.id,
        filename=chat_file.filename,
        mime_type=chat_file.mime_type,
        size=chat_file.size,
        file_type=chat_file.file_type,
    )


@router.get("/{file_id}", response_class=Response)
async def download_file(
{%- if cookiecutter.use_postgresql %}
    file_id: UUID,
{%- else %}
    file_id: str,
{%- endif %}
    db: DBSession = None,
    current_user: CurrentUser = None,
):
    """Download a file. Only the owner can access their files."""
    result = await db.execute(select(ChatFile).where(ChatFile.id == file_id))
    chat_file = result.scalar_one_or_none()

    if not chat_file:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    if chat_file.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

    storage = get_file_storage()
    data = await storage.load(chat_file.storage_path)

    return Response(
        content=data,
        media_type=chat_file.mime_type,
        headers={"Content-Disposition": f'inline; filename="{chat_file.filename}"'},
    )


@router.get("/{file_id}/info", response_model=FileInfo)
async def get_file_info(
{%- if cookiecutter.use_postgresql %}
    file_id: UUID,
{%- else %}
    file_id: str,
{%- endif %}
    db: DBSession = None,
    current_user: CurrentUser = None,
):
    """Get file metadata. Only the owner can access."""
    result = await db.execute(select(ChatFile).where(ChatFile.id == file_id))
    chat_file = result.scalar_one_or_none()

    if not chat_file:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
    if chat_file.user_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Access denied")

    return FileInfo(
        id=chat_file.id,
        filename=chat_file.filename,
        mime_type=chat_file.mime_type,
        size=chat_file.size,
        file_type=chat_file.file_type,
        created_at=chat_file.created_at,
        user_id=chat_file.user_id,
    )
{%- endif %}
