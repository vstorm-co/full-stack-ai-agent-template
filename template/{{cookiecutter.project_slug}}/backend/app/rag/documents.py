{%- if cookiecutter.enable_rag %}
from abc import ABC, abstractmethod
from pathlib import Path

from langchain_text_splitters import RecursiveCharacterTextSplitter
{%- if cookiecutter.use_llamaparse %}
from llama_cloud import AsyncLlamaCloud
{%- endif %}
import pdfplumber
from docx import Document as DOCXDocument

from app.rag.config import RAGSettings, DocumentExtensions
from app.rag.models import Document, DocumentMetadata, DocumentPage, DocumentPageChunk


class BaseDocumentParser(ABC):
    """Abstract base class for document parsing strategies.
    Defines the interface that all document parsers must implement.
    Supports parsing of various document formats (PDF, DOCX, TXT, MD).
    """
    
    allowed = [f"{ext.value}" for ext in DocumentExtensions]
    
    def is_file_existing(self, filepath: Path) -> bool:
        """Check if file exists at the given path.
        Args:
            filepath: Path to the file to check.
        Returns:
            True if the file exists, False otherwise.
        """
        return Path.exists(filepath)
    
    def is_extension_allowed(self, filepath: Path) -> bool:
        """Check whether document extension is allowed for parsing.
        Args:
            filepath: Path to the file to check.
        Returns:
            True if the extension is supported and file exists.
        """
        return filepath.suffix.lower() in self.allowed and self.is_file_existing(filepath)
    
    def get_document_metadata(self, filepath: Path) -> DocumentMetadata:
        """Collect metadata about a given document.
        Args:
            filepath: Path to the document file.
        Returns:
            DocumentMetadata object containing file information.
        """
        return DocumentMetadata(
            filename=filepath.name,
            filesize=filepath.stat().st_size,
            filetype=filepath.suffix.replace(".", "")
        )
    
    @abstractmethod
    def parse(self, filepath: Path) -> Document:
        """Parse a file and read its content into a Document object.
        Args:
            filepath: Path to the file to parse.
        Returns:
            Document object with parsed content and metadata.
        """
        pass


class TextDocumentParser(BaseDocumentParser):
    """Parser for text-based documents (TXT, MD).
    Uses Python's built-in file reading capabilities to extract
    text content from plain text and Markdown files.
    """
    
    def _parse_text_file(self, filepath: Path) -> Document:
        """Extract raw text from a TXT or MD file.
        Args:
            filepath: Path to the text file.
        Returns:
            Document object with the file content.
        """
        with open(filepath, "r", encoding="utf-8") as f:
            page = DocumentPage(
                page_num=1,
                content=f.read()
            )
            
        return Document(
            pages=[page], 
            metadata=self.get_document_metadata(filepath)
        )
    
    async def parse(self, filepath: Path) -> Document:
        """Parse a text file (TXT or MD).
        
        Args:
            filepath: Path to the text file.
            
        Returns:
            Document object with parsed content.
            
        Raises:
            ValueError: If the file extension is not supported.
        """
        if not self.is_extension_allowed(filepath):
            raise ValueError(f"Extension {filepath.suffix} not supported by TextDocumentParser")
        
        if filepath.suffix in (".txt", ".md"):
            return self._parse_text_file(filepath)
        else:
            raise ValueError(f"Unsupported file extension. Allowed extensions: {self.allowed}")


class DocxDocumentParser(BaseDocumentParser):
    """Parser for DOCX documents using python-docx.
    
    Extracts text content from Microsoft Word documents by reading
    all paragraphs and joining them with newline characters.
    """
    
    def _parse_docx_file(self, filepath: Path) -> Document:
        """Extract raw text from the DOCX file.
        
        Args:
            filepath: Path to the DOCX file.
            
        Returns:
            Document object with the file content.
        """
        file = DOCXDocument(filepath)
        page = DocumentPage(
            page_num=1,
            content="\n".join([p.text for p in file.paragraphs])
        )
        return Document(
            pages=[page],
            metadata=self.get_document_metadata(filepath)
        )
    
    async def parse(self, filepath: Path) -> Document:
        """Parse a DOCX file.
        
        Args:
            filepath: Path to the DOCX file.
            
        Returns:
            Document object with parsed content.
            
        Raises:
            ValueError: If the file is not a DOCX file.
        """
        if not self.is_extension_allowed(filepath):
            raise ValueError(f"Extension {filepath.suffix} not supported by DocxDocumentParser")
        
        if filepath.suffix == ".docx":
            return self._parse_docx_file(filepath)
        else:
            raise ValueError(f"Unsupported file extension. Allowed extensions: {self.allowed}")


class PdfPlumberParser(BaseDocumentParser):
    """Local PDF parser using pdfplumber.
    
    Extracts text from PDF files using the pdfplumber library.
    Note: Files that do not have a text layer will be treated as empty.
    """
    
    def _parse_pdf_file(self, filepath: Path) -> Document:
        """Extract raw text from a PDF file.
        
        Args:
            filepath: Path to the PDF file.
            
        Returns:
            Document object with pages containing extracted text.
            
        Note:
            Files that do not have a text layer are treated as empty.
        """
        with pdfplumber.open(filepath) as pdf:
            pages = []
            
            for page in pdf.pages:
                pages.append(
                    DocumentPage(
                        page_num=page.page_number,
                        content=page.extract_text() or ""
                    )
                )
                
        return Document(
            pages=pages,
            metadata=self.get_document_metadata(filepath)
        )
    
    async def parse(self, filepath: Path) -> Document:
        """Parse a PDF file.
        
        Args:
            filepath: Path to the PDF file.
            
        Returns:
            Document object with parsed content.
            
        Raises:
            ValueError: If the file is not a PDF file.
        """
        if not self.is_extension_allowed(filepath):
            raise ValueError(f"Extension {filepath.suffix} not supported by PdfPlumberParser")
        
        if filepath.suffix == ".pdf":
            return self._parse_pdf_file(filepath)
        else:
            raise ValueError(f"Unsupported file extension. Allowed extensions: {self.allowed}")


{% if cookiecutter.use_llamaparse -%}
class LlamaParseParser(BaseDocumentParser):
    """Advanced PDF parser using LlamaParse cloud API.
    
    Provides superior PDF parsing by using LlamaParse's AI-powered
    document understanding capabilities. Returns markdown-formatted content.
    """

    def __init__(self, api_key: str):
        """Initialize the LlamaParse parser.
        
        Args:
            api_key: LlamaCloud API key for authentication.
        """
        self.parser = AsyncLlamaCloud(api_key=api_key)

    async def parse(self, filepath: Path) -> Document:
        """Parse a PDF file using LlamaParse.
        
        Args:
            filepath: Path to the PDF file.
            
        Returns:
            Document object with parsed markdown content.
            
        Raises:
            ValueError: If the file is not a PDF file.
        """
        if not self.is_extension_allowed(filepath):
            raise ValueError(f"Extension {filepath.suffix} not supported by LlamaParse")
        
        if filepath.suffix != ".pdf":
            raise ValueError("LlamaParse is only supported for PDF files")
        
        # Upload and parse a document
        file_obj = await self.parser.files.create(file=filepath, purpose="parse")
        result = await self.parser.parsing.parse(
            file_id=file_obj.id,
            tier="agentic",
            version="latest",
            expand=["text", "markdown"],
        )
        pages = []
        for page in result.markdown.pages:
            pages.append(DocumentPage(
                page_num=page.page_number,
                content=page.markdown
            ))
            
        return Document(
            pages=pages,
            metadata=self.get_document_metadata(filepath)
        )
{%- endif %}


class DocumentProcessor:
    """Orchestrates parsing and chunking of files into Document objects.
    
    Manages the document processing pipeline:
    1. Route to appropriate parser based on file extension
    2. Parse document content
    3. Chunk document pages using RecursiveCharacterTextSplitter
    
    Supported file types:
    - TXT, MD: TextDocumentParser (Python native)
    - DOCX: DocxDocumentParser (Python native)
    - PDF: PdfPlumberParser or LlamaParseParser (based on configuration)
    """

    def __init__(self, settings: RAGSettings):
        """Initialize the document processor.
        
        Args:
            settings: RAG configuration settings.
        """
        self.settings = settings
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            is_separator_regex=False,
        )
        
        # Always use Python native parsers for non-PDF files
        self.text_parser = TextDocumentParser()
        self.docx_parser = DocxDocumentParser()
        
        # PDF parser based on configuration
        {%- if cookiecutter.use_llamaparse %}
        self.pdf_parser = LlamaParseParser(api_key=settings.pdf_parser.api_key)
        {%- else %}
        self.pdf_parser = PdfPlumberParser()
        {%- endif %}

    async def process_file(self, filepath: Path) -> Document:
        """Main entry point: filepath -> Document with chunks.
        
        Args:
            filepath: Path to the file to process.
            
        Returns:
            Document object with parsed pages and chunked content.
            
        Raises:
            ValueError: If the file type is not supported.
        """
        # Route to appropriate parser based on file extension
        if filepath.suffix in (".txt", ".md"):
            document = await self.text_parser.parse(filepath)
        elif filepath.suffix == ".docx":
            document = await self.docx_parser.parse(filepath)
        elif filepath.suffix == ".pdf":
            document = await self.pdf_parser.parse(filepath)
        else:
            raise ValueError(f"Unsupported file type: {filepath.suffix}")
        
        pages = document.pages

        chunked_pages: list[DocumentPageChunk] = []
        for page in pages:
            chunked_text = self.splitter.split_text(page.content)
            for chunk in chunked_text:
                chunked_pages.append(DocumentPageChunk(
                    chunk_content=chunk,
                    parent_doc_id=document.id,
                    **page.model_dump(
                        exclude={"parent_doc_id"}
                    )))
        
        # Add chunked pages to original document
        document.chunked_pages = chunked_pages
        return document

{%- endif %}
